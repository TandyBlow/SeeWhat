"""
Word .docx parsing.
"""


def parse_docx(file_path: str) -> str:
    """Parse Word .docx file and extract text from all paragraphs."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX parsing. Install it with: pip install python-docx"
        )

    doc = Document(file_path)
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text
        if text.strip():
            # Preserve heading styles with markdown-like prefix
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
                paragraphs.append('#' * min(level, 6) + ' ' + text)
            else:
                paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = '\t'.join(cell.text for cell in row.cells)
            if row_text.strip():
                paragraphs.append(row_text)

    return '\n\n'.join(paragraphs)
